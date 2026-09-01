"""
SHACK ENTERTAINMENT — shack_statement_drafter.py
The 7-day duty, automated. Reads Data\artist_sales.csv + Data\artists.csv,
fills the Artist Statement bones, banks one docx per artist with activity.
Usage: py shack_statement_drafter.py [YYYY-MM]   (default: last month)
"""
import os, sys, csv
from datetime import date, datetime
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

script_dir = os.path.dirname(os.path.abspath(__file__))
project_root = os.path.dirname(script_dir)
DATA = os.path.join(project_root, 'Data')
BRAND = os.path.join(project_root, 'Brand')
OUT = os.path.join(project_root, 'Statements')
os.makedirs(OUT, exist_ok=True)
NAVY, AMBER, GREY = '14263F', 'C5912A', '555555'

SALES_HEADER = 'month,artist_id,item,outlet,gross,commission'
SALES_PATH = os.path.join(DATA, 'artist_sales.csv')
if not os.path.exists(SALES_PATH):
    with open(SALES_PATH, 'w', encoding='utf-8') as f:
        f.write(SALES_HEADER + '\n')

LOGOS = [('logo_shack.png', 2.2), ('logo_au.png', 1.3),
         ('logo_le.png', 1.3), ('logo_snn.png', 1.3)]

def style_run(r, size, bold, color):
    r.font.size = Pt(size); r.font.bold = bold
    r.font.color.rgb = RGBColor.from_string(color); r.font.name = 'Arial'

def base():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.8); sec.right_margin = Cm(2.2)
    pgB = OxmlElement('w:pgBorders'); pgB.set(qn('w:offsetFrom'), 'page')
    lb = OxmlElement('w:left')
    for k, v in (('w:val', 'single'), ('w:sz', '24'), ('w:space', '18'), ('w:color', AMBER)):
        lb.set(qn(k), v)
    pgB.append(lb); sec._sectPr.append(pgB)

    def para(text, size, bold=False, color=NAVY, after=0):
        p = doc.add_paragraph()
        style_run(p.add_run(text), size, bold, color)
        p.paragraph_format.space_after = Pt(after)
        p.paragraph_format.space_before = Pt(0)
        return p

    tbl = doc.add_table(rows=1, cols=2)
    left, right = tbl.rows[0].cells
    p = left.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    style_run(p.add_run('SHACK'), 26, True, NAVY)
    p = left.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    style_run(p.add_run('ENTERTAINMENT'), 26, True, NAVY)
    added = 0
    for fn, w in LOGOS:
        path = os.path.join(BRAND, fn)
        if not os.path.exists(path):
            continue
        p = right.paragraphs[0] if added == 0 else right.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.space_before = Pt(0)
        p.add_run().add_picture(path, width=Cm(w))
        added += 1
    rule = doc.add_paragraph(); rule.paragraph_format.space_after = Pt(3)
    pBdr = OxmlElement('w:pBdr'); bt = OxmlElement('w:bottom')
    for k, v in (('w:val', 'single'), ('w:sz', '6'), ('w:space', '1'), ('w:color', NAVY)):
        bt.set(qn(k), v)
    pBdr.append(bt); rule._p.get_or_add_pPr().append(pBdr)
    para('ARTISTS UNLIMITED  ·  THE LIVE EXCHANGE  ·  SHACK NEWS NETWORK',
         9, color=GREY, after=18)
    fp = sec.footer.paragraphs[0]
    fr = fp.add_run('Shack Entertainment Limited · Company No. 14628241 · '
                    '25 Fielding Avenue, Twickenham, England TW2 5LX')
    fr.font.size = Pt(8); fr.font.color.rgb = RGBColor.from_string(GREY)
    fr.font.name = 'Arial'
    return doc, para

def month_arg():
    if len(sys.argv) > 1:
        return sys.argv[1]
    d = date.today().replace(day=1)
    m = d.month - 1 or 12
    y = d.year if d.month > 1 else d.year - 1
    return f"{y}-{m:02d}"

month = month_arg()
month_name = datetime.strptime(month + '-01', '%Y-%m-%d').strftime('%B %Y')

artists = {}
ap = os.path.join(DATA, 'artists.csv')
if os.path.exists(ap):
    with open(ap, encoding='utf-8') as f:
        for row in csv.DictReader(f):
            artists[(row.get('id') or '').strip()] = (row.get('name') or '').strip()

sales = {}
with open(SALES_PATH, encoding='utf-8') as f:
    for row in csv.DictReader(f):
        if (row.get('month') or '').strip() == month:
            sales.setdefault((row.get('artist_id') or '').strip(), []).append(row)

if not sales:
    print(f"No statements due for {month} — the honest zero.")
    raise SystemExit(0)

for aid, rows in sorted(sales.items()):
    name = artists.get(aid, 'Canary Artist')
    doc, para = base()
    para(f"ARTIST STATEMENT — {month_name.upper()}", 14, True, after=4)
    para(f"{name} — {aid}", 11, after=12)
    t = doc.add_table(rows=1 + len(rows), cols=5)
    t.style = 'Table Grid'
    for i, h in enumerate(['Item', 'Outlet', 'Gross', 'Commission', 'Net']):
        t.rows[0].cells[i].text = h
    tot_g = tot_c = 0.0
    for i, r in enumerate(rows, start=1):
        g = float(r['gross']); c = float(r['commission'])
        tot_g += g; tot_c += c
        t.rows[i].cells[0].text = r['item']
        t.rows[i].cells[1].text = r['outlet']
        t.rows[i].cells[2].text = f"{g:.2f}"
        t.rows[i].cells[3].text = f"{c:.2f}"
        t.rows[i].cells[4].text = f"{g - c:.2f}"
    para('', 6)
    para(f"Artist Net Payable: £{tot_g - tot_c:.2f}", 12, True, after=4)
    para('Paid within 7 days of month-end from the Santander Client Account.',
         9, color=GREY)
    fn = f"Statement_{aid}_{month}.docx"
    doc.save(os.path.join(OUT, fn))
    print('Banked:', fn)