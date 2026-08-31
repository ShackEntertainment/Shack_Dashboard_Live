"""
SHACK ENTERTAINMENT — shack_generate_stationery.py
Production stationery family: master letterhead, order form,
artist statement — one spine, typed crisp.
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

script_dir = os.path.dirname(os.path.abspath(__file__))
BRAND = os.path.join(os.path.dirname(script_dir), 'Brand')
os.makedirs(BRAND, exist_ok=True)
NAVY, AMBER, GREY = '14263F', 'C5912A', '555555'

def base():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.8); sec.right_margin = Cm(2.2)
    pgB = OxmlElement('w:pgBorders'); pgB.set(qn('w:offsetFrom'), 'page')
    lb = OxmlElement('w:left')
    for k, v in (('w:val', 'single'), ('w:sz', '24'), ('w:space', '18'), ('w:color', AMBER)):
        lb.set(qn(k), v)
    pgB.append(lb); sec._sectPr.append(pgB)

    def para(text, size, bold=False, color=NAVY, after=0):
        p = doc.add_paragraph(); r = p.add_run(text)
        r.font.size = Pt(size); r.font.bold = bold
        r.font.color.rgb = RGBColor.from_string(color); r.font.name = 'Arial'
        p.paragraph_format.space_after = Pt(after)
        p.paragraph_format.space_before = Pt(0)
        return p

    para('SHACK', 26, True); para('ENTERTAINMENT', 26, True, after=2)
    rule = doc.add_paragraph(); rule.paragraph_format.space_after = Pt(3)
    pBdr = OxmlElement('w:pBdr'); bt = OxmlElement('w:bottom')
    for k, v in (('w:val', 'single'), ('w:sz', '6'), ('w:space', '1'), ('w:color', NAVY)):
        bt.set(qn(k), v)
    pBdr.append(bt); rule._p.get_or_add_pPr().append(pBdr)
    para('ARTISTS UNLIMITED  ·  THE LIVE EXCHANGE  ·  SHACK NEWS NETWORK',
         9, color=GREY, after=18)
    fp = sec.footer.paragraphs[0]
    fr = fp.add_run('Shack Entertainment Limited · Company No. 14628241 · '
                    '25 Fielding Avenue, Twickenham, England TW2 5LX')
    fr.font.size = Pt(8); fr.font.color.rgb = RGBColor.from_string(GREY)
    fr.font.name = 'Arial'
    return doc, para

# 1 — master letterhead
doc, para = base()
doc.add_paragraph()
doc.save(os.path.join(BRAND, 'Shack_Letterhead_Master.docx'))
print('Banked: Shack_Letterhead_Master.docx')

# 2 — order form (LE_02 / SNN_02 print path)
doc, para = base()
para('ENGAGEMENT & COMMISSION ORDER', 14, True, after=12)
tbl = doc.add_table(rows=11, cols=2)
tbl.style = 'Table Grid'
fields = ['Order Ref', 'Name', 'ID (LE / SNN / ART)', 'Role / Discipline',
          'Event or Piece', 'Venue or Outlet', 'Date(s)', 'Fee Type',
          'Fee / Split', 'Backline / Rider', 'Notes']
for i, f in enumerate(fields):
    tbl.rows[i].cells[0].text = f
para('', 6)
para('For the Company: ………………………………     For the Talent: ………………………………', 10, after=4)
doc.save(os.path.join(BRAND, 'Shack_Order_Form.docx'))
print('Banked: Shack_Order_Form.docx')

# 3 — artist statement (the 7-day duty's bones)
doc, para = base()
para('ARTIST STATEMENT — [MONTH YEAR]', 14, True, after=4)
para('[Artist Name] — [ID]', 11, after=12)
tbl = doc.add_table(rows=6, cols=5)
tbl.style = 'Table Grid'
for i, h in enumerate(['Item', 'Outlet', 'Gross', 'Commission', 'Net']):
    tbl.rows[0].cells[i].text = h
para('', 6)
para('Artist Net Payable: £…………', 12, True, after=4)
para('Paid within 7 days of month-end from the Santander Client Account.', 9, color=GREY)
doc.save(os.path.join(BRAND, 'Shack_Artist_Statement.docx'))
print('Banked: Shack_Artist_Statement.docx')