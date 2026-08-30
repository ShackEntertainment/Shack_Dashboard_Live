"""
SHACK ENTERTAINMENT — shack_generate_letterhead.py
Builds the production letterhead (.docx) from the Brand spec:
amber left rule, navy wordmark, lane line, typed legal footer.
Generated art was the spec; this file is the print path.
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

script_dir = os.path.dirname(os.path.abspath(__file__))
BRAND = os.path.join(os.path.dirname(script_dir), 'Brand')
os.makedirs(BRAND, exist_ok=True)

NAVY = '14263F'
AMBER = 'C5912A'
GREY = '555555'

doc = Document()
sec = doc.sections[0]
sec.top_margin = Cm(2.0)
sec.bottom_margin = Cm(2.0)
sec.left_margin = Cm(2.8)
sec.right_margin = Cm(2.2)

# Amber left rule — full-height page border
pgB = OxmlElement('w:pgBorders')
pgB.set(qn('w:offsetFrom'), 'page')
lb = OxmlElement('w:left')
lb.set(qn('w:val'), 'single')
lb.set(qn('w:sz'), '24')
lb.set(qn('w:space'), '18')
lb.set(qn('w:color'), AMBER)
pgB.append(lb)
sec._sectPr.append(pgB)

def para(text, size, bold=False, color=NAVY, after=0):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.color.rgb = RGBColor.from_string(color)
    r.font.name = 'Arial'
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(0)
    return p

# Wordmark
para('SHACK', 26, bold=True)
para('ENTERTAINMENT', 26, bold=True, after=2)

# Navy rule under the wordmark
rule = doc.add_paragraph()
rule.paragraph_format.space_after = Pt(3)
pBdr = OxmlElement('w:pBdr')
bt = OxmlElement('w:bottom')
bt.set(qn('w:val'), 'single')
bt.set(qn('w:sz'), '6')
bt.set(qn('w:space'), '1')
bt.set(qn('w:color'), NAVY)
pBdr.append(bt)
rule._p.get_or_add_pPr().append(pBdr)

# Lane line — formal names, per the brand rule
para('ARTISTS UNLIMITED  ·  THE LIVE EXCHANGE  ·  SHACK NEWS NETWORK',
     9, color=GREY, after=24)

doc.add_paragraph()

# Legal footer — typed verbatim, on every page
fp = sec.footer.paragraphs[0]
fr = fp.add_run('Shack Entertainment Limited · Company No. 14628241 · '
                '25 Fielding Avenue, Twickenham, England TW2 5LX')
fr.font.size = Pt(8)
fr.font.color.rgb = RGBColor.from_string(GREY)
fr.font.name = 'Arial'

out = os.path.join(BRAND, 'Shack_Letterhead_Master.docx')
doc.save(out)
print('Banked:', out)