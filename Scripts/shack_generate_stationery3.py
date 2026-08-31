"""
SHACK ENTERTAINMENT — shack_generate_stationery3.py
Stationery v3 — clean masthead table, department marks stacked right.
Usage:
  py shack_generate_stationery3.py            equal marks
  py shack_generate_stationery3.py au|le|snn  that lane's mark enlarged
Logos: Brand\logo_shack.png / logo_au.png / logo_le.png / logo_snn.png
"""
import os, sys
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

script_dir = os.path.dirname(os.path.abspath(__file__))
BRAND = os.path.join(os.path.dirname(script_dir), 'Brand')
NAVY, AMBER, GREY = '14263F', 'C5912A', '555555'
FOCUS = sys.argv[1].lower() if len(sys.argv) > 1 else None
SUFFIX = ('_' + FOCUS.upper()) if FOCUS in ('au', 'le', 'snn') else ''

LOGOS = [('logo_shack.png', None, 2.2),
         ('logo_au.png', 'au', 1.3),
         ('logo_le.png', 'le', 1.3),
         ('logo_snn.png', 'snn', 1.3)]

def style_run(r, size, bold, color):
    r.font.size = Pt(size); r.font.bold = bold
    r.font.color.rgb = RGBColor.from_string(color); r.font.name = 'Arial'

def base():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.8); sec.right_margin = Cm(2.2)
    pgB = OxmlElement('w:pgBorders'); pgB.set(qn('w:offsetFrom'), 'page')
    lb = OxmlElement('w:left')
    for k, v in (('w:val', 'single'), ('w:sz', '24'), ('w:space', '18'), ('w:color', AMBER)):
        lb.set(qn(k), v)
    pgB.append(lb); sec._sectPr.append(pgB)

    def para(text, size, bold=False, color=NAVY, after=0):
        p = doc.add_paragraph()
        style_run(p.add_run(text), size, bold, color)
        p.paragraph_format.space_after = Pt(after)
        p.paragraph_format.space_before = Pt(0)
        return p

    tbl = doc.add_table(rows=1, cols=2)
    left, right = tbl.rows[0].cells
    left.width = Cm(12); right.width = Cm(5)
    p = left.paragraphs[0]
    p.paragraph_format.space_after = Pt(0); p.paragraph_format.space_before = Pt(0)
    style_run(p.add_run('SHACK'), 26, True, NAVY)
    p = left.add_paragraph()
    p.paragraph_format.space_after = Pt(2); p.paragraph_format.space_before = Pt(0)
    style_run(p.add_run('ENTERTAINMENT'), 26, True, NAVY)

    added = 0
    for fn, lane, w in LOGOS:
        path = os.path.join(BRAND, fn)
        if not os.path.exists(path):
            print('missing logo:', fn)
            continue
        p = right.paragraphs[0] if added == 0 else right.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.space_before = Pt(0)
        p.add_run().add_picture(path, width=Cm(1.9 if lane == FOCUS else w))
        added += 1
    if added == 0:
        print('WARNING: no logos found — text-only masthead.')

    rule = doc.add_paragraph(); rule.paragraph_format.space_after = Pt(3)
    pBdr = OxmlElement('w:pBdr'); bt = OxmlElement('w:bottom')
    for k, v in (('w:val', 'single'), ('w:sz', '6'), ('w:space', '1'), ('w:color', NAVY)):
        bt.set(qn(k), v)
    pBdr.append(bt); rule._p.get_or_add_pPr().append(pBdr)
    para('ARTISTS UNLIMITED  ·  THE LIVE EXCHANGE  ·  SHACK NEWS NETWORK',
         9, color=GREY, after=18)
    fp = sec.footer.paragraphs[0]
    fr = fp.add_run('Shack Entertainment Limited · Company No. 14628241 · '
                    '25 Fielding Avenue, Twickenham, England TW2 5LX')
    fr.font.size = Pt(8); fr.font.color.rgb = RGBColor.from_string(GREY)
    fr.font.name = 'Arial'
    return doc, para

doc, para = base()
doc.add_paragraph()
doc.save(os.path.join(BRAND, 'Shack_Letterhead_Master%s.docx' % SUFFIX))
print('Banked: Shack_Letterhead_Master%s.docx' % SUFFIX)

doc, para = base()
para('ENGAGEMENT & COMMISSION ORDER', 14, True, after=12)
t = doc.add_table(rows=11, cols=2); t.style = 'Table Grid'
for i, f in enumerate(['Order Ref', 'Name', 'ID (LE / SNN / ART)', 'Role / Discipline',
                       'Event or Piece', 'Venue or Outlet', 'Date(s)', 'Fee Type',
                       'Fee / Split', 'Backline / Rider', 'Notes']):
    t.rows[i].cells[0].text = f
para('', 6)
para('For the Company: ………………………………     For the Talent: ………………………………', 10, after=4)
doc.save(os.path.join(BRAND, 'Shack_Order_Form%s.docx' % SUFFIX))
print('Banked: Shack_Order_Form%s.docx' % SUFFIX)

doc, para = base()
para('ARTIST STATEMENT — [MONTH YEAR]', 14, True, after=4)
para('[Artist Name] — [ID]', 11, after=12)
t = doc.add_table(rows=6, cols=5); t.style = 'Table Grid'
for i, h in enumerate(['Item', 'Outlet', 'Gross', 'Commission', 'Net']):
    t.rows[0].cells[i].text = h
para('', 6)
para('Artist Net Payable: £…………', 12, True, after=4)
para('Paid within 7 days of month-end from the Santander Client Account.', 9, color=GREY)
doc.save(os.path.join(BRAND, 'Shack_Artist_Statement%s.docx' % SUFFIX))
print('Banked: Shack_Artist_Statement%s.docx' % SUFFIX)